"""Minimal markdown -> docx converter using python-docx.

Why this exists: pandoc-generated .docx files frequently trigger Word's
"file is corrupt" / protected-view validator on Windows. python-docx writes
OOXML in the exact shape Word expects, which avoids the issue. Use this
script for any markdown -> .docx conversion in this repo; do not rely on
`pandoc -o file.docx` for files that will be opened in Word.

Supports the subset used by docs/rapporten/*.md:
  - ATX headings (# .. ###)
  - Paragraphs
  - Bullet lists (-, *)
  - Numbered lists (1. 2. ...)
  - GitHub-flavored pipe tables
  - Inline **bold** and *italic*
  - Horizontal rule (---)

Usage:
    python scripts/md_to_docx.py INPUT.md [OUTPUT.docx]
    python scripts/md_to_docx.py INPUT.md INPUT2.md ...   (writes .docx alongside)
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*)")


def add_inline_runs(paragraph, text: str) -> None:
    for piece in INLINE_RE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) >= 4:
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) >= 2:
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        else:
            paragraph.add_run(piece)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.match(r"^:?-+:?$", c.strip()) for c in rows[1]):
        header = rows[0]
        body = rows[2:]
        return [header] + body, i
    return rows, i


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("---") and set(stripped) == {"-"}:
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            content = heading_match.group(2).strip()
            heading = doc.add_heading(level=min(level, 4))
            add_inline_runs(heading, content)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].lstrip().startswith("|"):
            rows, new_i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        add_inline_runs(para, cell_text)
                        if r_idx == 0:
                            for run in para.runs:
                                run.bold = True
                doc.add_paragraph()
                i = new_i
                continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            para = doc.add_paragraph(style="List Bullet")
            add_inline_runs(para, bullet_match.group(1))
            i += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            para = doc.add_paragraph(style="List Number")
            add_inline_runs(para, numbered_match.group(1))
            i += 1
            continue

        # Collect soft-wrapped paragraph lines until blank, heading, list, or table.
        buf = [stripped]
        j = i + 1
        while j < len(lines):
            peek = lines[j].strip()
            if not peek:
                break
            if re.match(r"^#{1,6}\s", peek):
                break
            if re.match(r"^[-*]\s", peek) or re.match(r"^\d+\.\s", peek):
                break
            if peek.startswith("|"):
                break
            buf.append(peek)
            j += 1
        para = doc.add_paragraph()
        add_inline_runs(para, " ".join(buf))
        i = j

    doc.save(out_path)


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print(__doc__, file=sys.stderr)
        return 2
    args = argv[1:]
    if len(args) == 2 and args[1].lower().endswith(".docx"):
        convert(Path(args[0]), Path(args[1]))
        print(f"wrote {args[1]}")
        return 0
    for src in args:
        src_path = Path(src)
        out = src_path.with_suffix(".docx")
        convert(src_path, out)
        print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
